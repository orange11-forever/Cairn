import warnings
from io import BytesIO
from itertools import pairwise
from typing import Any, cast
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from cairn_worker.errors import WorkerFailure
from cairn_worker.parsers import BlockKind
from cairn_worker.parsers.docx import DocxParser
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE

from apps.worker.tests.fixture_factory import empty_docx_fixture


def _save_docx(document: DocumentObject) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _structured_docx() -> bytes:
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.add_run("Visible ")
    hidden = paragraph.add_run("private hidden ")
    hidden.font.hidden = True
    paragraph.add_run("世界")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "一"
    document.add_heading("Details", level=2)
    document.add_paragraph("Tail")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _append_member(package: bytes, name: str, content: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(package), "r") as source, ZipFile(
        output, "w", compression=ZIP_DEFLATED
    ) as target:
        for member in source.infolist():
            target.writestr(member, source.read(member))
        target.writestr(name, content)
    return output.getvalue()


def _replace_member(package: bytes, name: str, content: bytes) -> bytes:
    return _replace_members(package, {name: content})


def _replace_members(package: bytes, replacements: dict[str, bytes]) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(package), "r") as source, ZipFile(
        output, "w", compression=ZIP_DEFLATED
    ) as target:
        for member in source.infolist():
            target.writestr(
                member,
                replacements.get(member.filename, source.read(member)),
            )
    return output.getvalue()


def _duplicate_first_member(package: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(package), "r") as source, ZipFile(
        output, "w", compression=ZIP_DEFLATED
    ) as target:
        members = source.infolist()
        for member in members:
            target.writestr(member, source.read(member))
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            target.writestr(members[0].filename, source.read(members[0]))
    return output.getvalue()


def _external_hyperlink_docx() -> bytes:
    document = Document()
    paragraph = document.add_paragraph("Before ")
    paragraph.add_run("Visible link label")
    paragraph.add_run(" after")
    output = BytesIO()
    document.save(output)
    package = output.getvalue()
    with ZipFile(BytesIO(package), "r") as source:
        document_xml = source.read("word/document.xml")
        relationships = source.read("word/_rels/document.xml.rels")
    visible_run = b"<w:r><w:t>Visible link label</w:t></w:r>"
    assert visible_run in document_xml
    document_xml = document_xml.replace(
        visible_run,
        b'<w:hyperlink r:id="rIdExternal"><w:r><w:t>Visible link label</w:t></w:r>'
        b"</w:hyperlink>",
        1,
    )
    relationships = relationships.replace(
        b"</Relationships>",
        b'<Relationship Id="rIdExternal" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
        b'Target="https://127.0.0.1:1/private" TargetMode="External"/></Relationships>',
        1,
    )
    return _replace_members(
        package,
        {
            "word/document.xml": document_xml,
            "word/_rels/document.xml.rels": relationships,
        },
    )


def _mark_first_member_encrypted(package: bytes) -> bytes:
    marked = bytearray(package)
    local = marked.find(b"PK\x03\x04")
    central = marked.find(b"PK\x01\x02")
    assert local >= 0 and central >= 0
    local_flags = int.from_bytes(marked[local + 6 : local + 8], "little") | 1
    central_flags = int.from_bytes(marked[central + 8 : central + 10], "little") | 1
    marked[local + 6 : local + 8] = local_flags.to_bytes(2, "little")
    marked[central + 8 : central + 10] = central_flags.to_bytes(2, "little")
    return bytes(marked)


def _corrupt_member_payload(package: bytes, name: str) -> bytes:
    damaged = bytearray(package)
    with ZipFile(BytesIO(package)) as source:
        member = source.getinfo(name)
        offset = member.header_offset
        filename_length = int.from_bytes(damaged[offset + 26 : offset + 28], "little")
        extra_length = int.from_bytes(damaged[offset + 28 : offset + 30], "little")
        data_offset = offset + 30 + filename_length + extra_length
        damaged[data_offset] ^= 0x80
    return bytes(damaged)


def _overlap_member_payload(package: bytes, name: str) -> bytes:
    damaged = bytearray(package)
    cursor = 0
    encoded_name = name.encode()
    while True:
        central = damaged.find(b"PK\x01\x02", cursor)
        assert central >= 0
        filename_length = int.from_bytes(damaged[central + 28 : central + 30], "little")
        filename = bytes(damaged[central + 46 : central + 46 + filename_length])
        if filename == encoded_name:
            compressed_size = int.from_bytes(damaged[central + 20 : central + 24], "little")
            damaged[central + 20 : central + 24] = (compressed_size + 1_024).to_bytes(4, "little")
            return bytes(damaged)
        cursor = central + 46 + filename_length


def _revision_docx() -> bytes:
    document = Document()
    document.add_heading("Heading", level=1)
    document.add_paragraph("Before accepted removed After")
    output = BytesIO()
    document.save(output)
    with ZipFile(BytesIO(output.getvalue())) as package:
        xml = package.read("word/document.xml")
    xml = xml.replace(
        b"<w:r><w:t>Before accepted removed After</w:t></w:r>",
        b"<w:r><w:t xml:space=\"preserve\">Before </w:t></w:r>"
        b"<w:ins><w:r><w:t>accepted</w:t></w:r></w:ins>"
        b"<w:del><w:r><w:delText> removed</w:delText></w:r></w:del>"
        b"<w:r><w:t xml:space=\"preserve\"> After</w:t></w:r>",
    )
    return _replace_member(output.getvalue(), "word/document.xml", xml)


def _block_revision_docx() -> bytes:
    document = Document()
    document.add_paragraph("before")
    document.add_paragraph("inserted block")
    document.add_paragraph("deleted block")
    document.add_paragraph("after")
    package = _save_docx(document)
    with ZipFile(BytesIO(package)) as source:
        xml = source.read("word/document.xml")
    for text, wrapper in ((b"inserted block", b"ins"), (b"deleted block", b"del")):
        marker = b"<w:t>" + text + b"</w:t>"
        marker_position = xml.index(marker)
        start = xml.rfind(b"<w:p", 0, marker_position)
        end = xml.index(b"</w:p>", marker_position) + len(b"</w:p>")
        paragraph = xml[start:end]
        xml = xml[:start] + b"<w:" + wrapper + b">" + paragraph + b"</w:" + wrapper + b">" + xml[end:]
    return _replace_member(package, "word/document.xml", xml)


def _nested_merged_hidden_docx() -> bytes:
    document = Document()
    document.add_heading("Heading", level=1)
    outer = document.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "outer"
    outer.cell(0, 1).text = "right"
    outer.cell(1, 0).merge(outer.cell(1, 1)).text = "horizontal"
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "nested"
    vertical = document.add_table(rows=2, cols=1)
    vertical.cell(0, 0).merge(vertical.cell(1, 0)).text = "vertical"
    hidden_character = document.styles.add_style(  # pyright: ignore[reportUnknownMemberType]
        "HiddenCharacter", WD_STYLE_TYPE.CHARACTER
    )
    hidden_character = cast(Any, hidden_character)
    hidden_character.font.hidden = True
    derived_character = document.styles.add_style(  # pyright: ignore[reportUnknownMemberType]
        "DerivedHidden", WD_STYLE_TYPE.CHARACTER
    )
    derived_character = cast(Any, derived_character)
    derived_character.base_style = hidden_character
    hidden_paragraph = document.styles.add_style(  # pyright: ignore[reportUnknownMemberType]
        "HiddenParagraph", WD_STYLE_TYPE.PARAGRAPH
    )
    hidden_paragraph = cast(Any, hidden_paragraph)
    hidden_paragraph.font.hidden = True
    paragraph = document.add_paragraph()
    paragraph.add_run("visible ")
    inherited = paragraph.add_run("character-secret ")
    inherited.style = derived_character
    paragraph.add_run("tail")
    hidden = document.add_paragraph(style=hidden_paragraph)
    hidden.add_run("paragraph-secret")
    override = hidden.add_run(" override-visible")
    override.font.hidden = False
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_parser_preserves_body_order_heading_paths_and_exact_ordinals() -> None:
    """Break caught: body XML order and exact paragraph/table locators must stay interleaved."""
    blocks = DocxParser().parse(BytesIO(_structured_docx()))

    assert [block.kind for block in blocks] == [
        BlockKind.HEADING,
        BlockKind.PARAGRAPH,
        BlockKind.TABLE,
        BlockKind.HEADING,
        BlockKind.PARAGRAPH,
    ]
    assert [block.text for block in blocks] == [
        "Overview",
        "Visible 世界",
        "name\tvalue\nalpha\t一",
        "Details",
        "Tail",
    ]
    assert [block.locator.model_dump(by_alias=True) for block in blocks] == [
        {"type": "docx", "headingPath": ["Overview"], "paragraph": 1, "table": None},
        {"type": "docx", "headingPath": ["Overview"], "paragraph": 3, "table": None},
        {"type": "docx", "headingPath": ["Overview"], "paragraph": None, "table": 1},
        {
            "type": "docx",
            "headingPath": ["Overview", "Details"],
            "paragraph": 4,
            "table": None,
        },
        {
            "type": "docx",
            "headingPath": ["Overview", "Details"],
            "paragraph": 5,
            "table": None,
        },
    ]
    locators = [block.locator.model_dump(by_alias=True) for block in blocks]
    assert all((locator["paragraph"] is None) != (locator["table"] is None) for locator in locators)


def test_docx_parser_rejects_an_empty_document() -> None:
    """Break caught: the default empty Word paragraph must not become searchable content."""
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(empty_docx_fixture()))

    assert caught.value.code == "no_extractable_text"


def test_docx_parser_indexes_only_an_external_hyperlinks_visible_label() -> None:
    """Break caught: Word relationship targets must never be fetched or included in text."""
    blocks = DocxParser().parse(BytesIO(_external_hyperlink_docx()))

    assert [block.text for block in blocks] == ["Before Visible link label after"]
    assert "127.0.0.1" not in blocks[0].text


@pytest.mark.parametrize(
    "content",
    [b"not a zip", _replace_member(empty_docx_fixture(), "word/document.xml", b"<broken")],
    ids=("not-opc", "malformed-document-xml"),
)
def test_docx_parser_rejects_malformed_packages_safely(content: bytes) -> None:
    """Break caught: ZIP/XML/library diagnostics must remain a bounded permanent fact."""
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(content))

    assert caught.value.code == "parser_failed"
    assert caught.value.retryable is False
    assert caught.value.safe_detail == "worker handler or parser failed"


@pytest.mark.parametrize(
    "content",
    [
        _append_member(empty_docx_fixture(), "word/vbaProject.bin", b"private macro"),
        _mark_first_member_encrypted(empty_docx_fixture()),
        _append_member(empty_docx_fixture(), "word/bomb.bin", b"0" * (1024 * 1024)),
        _append_member(empty_docx_fixture(), "../private.xml", b"unsafe"),
    ],
    ids=("macro", "encrypted-entry", "compression-ratio", "unsafe-path"),
)
def test_docx_parser_rejects_hazardous_opc_packages(content: bytes) -> None:
    """Break caught: Office libraries must not open macro, encrypted, bomb, or unsafe packages."""
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(content))

    assert caught.value.code == "parser_failed"
    assert caught.value.retryable is False
    assert caught.value.safe_detail == "worker handler or parser failed"


@pytest.mark.parametrize(
    "content",
    [
        _append_member(
            _append_member(empty_docx_fixture(), "EncryptionInfo", b"agile metadata"),
            "EncryptedPackage",
            b"encrypted payload",
        ),
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1EncryptionInfo\x00EncryptedPackage",
        _duplicate_first_member(empty_docx_fixture()),
    ],
    ids=("hybrid-encryption-streams", "ole-encrypted-container", "duplicate-name"),
)
def test_docx_parser_rejects_additional_encryption_and_duplicate_package_signals(
    content: bytes,
) -> None:
    """Break caught: non-flag encryption containers and duplicate names must fail pre-library."""
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(content))

    assert caught.value.code == "parser_failed"
    assert caught.value.retryable is False


def test_docx_parser_rejects_xml_doctype_and_external_entity_declarations() -> None:
    """Break caught: Office XML must never resolve an uploaded local or external entity."""
    package = empty_docx_fixture()
    with ZipFile(BytesIO(package), "r") as source:
        document_xml = source.read("word/document.xml")
    hostile_xml = document_xml.replace(
        b"<w:document",
        b'<!DOCTYPE w:document [<!ENTITY xxe SYSTEM "file:///private">]><w:document',
        1,
    ).replace(b"<w:body>", b"<w:body><w:p><w:r><w:t>&xxe;</w:t></w:r></w:p>", 1)

    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_replace_member(package, "word/document.xml", hostile_xml)))

    assert caught.value.code == "parser_failed"
    assert caught.value.retryable is False
    assert "private" not in caught.value.safe_detail


def test_docx_parser_enforces_opc_entry_and_uncompressed_size_limits(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Break caught: central-directory entry and expansion totals must be checked pre-library."""
    from cairn_worker.parsers import office_safety

    content = empty_docx_fixture()
    monkeypatch.setattr(office_safety, "OPC_MAX_ENTRIES", 1)
    with pytest.raises(WorkerFailure) as entries:
        DocxParser().parse(BytesIO(content))
    assert entries.value.code == "parser_failed"

    monkeypatch.setattr(office_safety, "OPC_MAX_ENTRIES", 10_000)
    monkeypatch.setattr(office_safety, "OPC_MAX_UNCOMPRESSED_BYTES", 1)
    with pytest.raises(WorkerFailure) as expansion:
        DocxParser().parse(BytesIO(content))
    assert expansion.value.code == "parser_failed"


def test_docx_parser_handles_revisions_nested_tables_merges_and_inherited_hidden_text() -> None:
    revision_blocks = DocxParser().parse(BytesIO(_revision_docx()))
    assert [block.text for block in revision_blocks] == ["Heading", "Before accepted After"]

    blocks = DocxParser().parse(BytesIO(_nested_merged_hidden_docx()))
    assert [block.text for block in blocks] == [
        "Heading",
        "outer\nnested\tright\nhorizontal",
        "vertical",
        "visible tail",
        "override-visible",
    ]
    assert [block.locator.model_dump(by_alias=True) for block in blocks[1:3]] == [
        {"type": "docx", "headingPath": ["Heading"], "paragraph": None, "table": 1},
        {"type": "docx", "headingPath": ["Heading"], "paragraph": None, "table": 2},
    ]
    assert blocks[1].text.count("nested") == 1

    block_revision = DocxParser().parse(BytesIO(_block_revision_docx()))
    assert [block.text for block in block_revision] == ["before", "inserted block", "after"]
    assert [block.locator.model_dump()["paragraph"] for block in block_revision] == [1, 2, 3]


@pytest.mark.parametrize("encoding", ["utf-8", "utf-16-le", "utf-16-be"])
@pytest.mark.parametrize("with_bom", [False, True])
def test_docx_parser_rejects_encoded_xml_entity_declarations(
    encoding: str,
    with_bom: bool,
) -> None:
    package = empty_docx_fixture()
    hostile = '<?xml version="1.0" encoding="%s"?><!DOCTYPE x [<!ENTITY xxe "x">]><x/>' % (
        "UTF-8" if encoding == "utf-8" else encoding.upper()
    )
    encoded = hostile.encode(encoding)
    if with_bom:
        encoded = {
            "utf-8": b"\xef\xbb\xbf",
            "utf-16-le": b"\xff\xfe",
            "utf-16-be": b"\xfe\xff",
        }[encoding] + encoded
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_replace_member(package, "word/document.xml", encoded)))
    assert caught.value.code == "parser_failed"


def test_docx_parser_rejects_utf16_relationship_entity_declaration() -> None:
    hostile = (
        '<?xml version="1.0" encoding="UTF-16BE"?>'
        '<!DOCTYPE Relationships [<!ENTITY xxe "private">]><Relationships>&xxe;</Relationships>'
    ).encode("utf-16-be")
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(
            BytesIO(
                _replace_member(
                    empty_docx_fixture(),
                    "word/_rels/document.xml.rels",
                    hostile,
                )
            )
        )
    assert caught.value.code == "parser_failed"


@pytest.mark.parametrize("encoding", ["utf-16-le", "utf-16-be"])
def test_docx_parser_rejects_encoded_macro_content_type_markers(encoding: str) -> None:
    package = empty_docx_fixture()
    with ZipFile(BytesIO(package)) as source:
        content_types = source.read("[Content_Types].xml").decode("utf-8")
    content_types = content_types.replace(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
        "application/vnd.ms-word.document.macroEnabled.main+xml",
    )
    declaration = f'<?xml version="1.0" encoding="{encoding.upper()}"?>'
    content_types = content_types[content_types.find("?>") + 2 :]
    encoded = (declaration + content_types).encode(encoding)
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_replace_member(package, "[Content_Types].xml", encoded)))
    assert caught.value.code == "parser_failed"


@pytest.mark.parametrize(
    "declaration",
    [
        (
            b'<Override PartName="/custom/arbitrary.bin" '
            b'ContentType="application/vnd.ms-word.document.macroEnabl&#101;d.main+xml"/>'
        ),
        (
            b'<Default Extension="bin" '
            b'ContentType="application/vnd.ms-office.vbaprojec&#116;"/>'
        ),
    ],
)
def test_docx_parser_rejects_numeric_reference_macro_content_type_before_library(
    monkeypatch: pytest.MonkeyPatch,
    declaration: bytes,
) -> None:
    """Break caught: XML references must not conceal semantic macro content types."""
    from cairn_worker.parsers import docx as docx_parser

    package = empty_docx_fixture()
    with ZipFile(BytesIO(package)) as source:
        content_types = source.read("[Content_Types].xml")
    content_types = content_types.replace(
        b"</Types>",
        declaration + b"</Types>",
    )
    opened = [False]

    def unexpected_document(*args: object, **kwargs: object) -> object:
        del args, kwargs
        opened[0] = True
        raise AssertionError("macro package reached python-docx")

    monkeypatch.setattr(docx_parser, "Document", unexpected_document)
    hostile = _append_member(
        _replace_member(package, "[Content_Types].xml", content_types),
        "custom/arbitrary.bin",
        b"macro",
    )
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(hostile))
    assert caught.value.code == "parser_failed"
    assert opened == [False]


def test_docx_parser_rejects_numeric_reference_internal_vba_relationship_type(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Break caught: arbitrary internal parts must not hide vbaProject relationship types."""
    from cairn_worker.parsers import docx as docx_parser

    package = empty_docx_fixture()
    with ZipFile(BytesIO(package)) as source:
        relationships = source.read("word/_rels/document.xml.rels")
    relationships = relationships.replace(
        b"</Relationships>",
        b'<Relationship Id="rMacro" '
        b'Type="http://schemas.example.test/relationships/vbaProjec&#116;" '
        b'Target="../custom/arbitrary.bin"/></Relationships>',
    )
    opened = [False]

    def unexpected_document(*args: object, **kwargs: object) -> object:
        del args, kwargs
        opened[0] = True
        raise AssertionError("macro relationship reached python-docx")

    monkeypatch.setattr(docx_parser, "Document", unexpected_document)
    hostile = _append_member(
        _replace_member(package, "word/_rels/document.xml.rels", relationships),
        "custom/arbitrary.bin",
        b"macro",
    )
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(hostile))
    assert caught.value.code == "parser_failed"
    assert opened == [False]


def test_docx_parser_allows_external_vba_relationship_label_without_fetch() -> None:
    package = empty_docx_fixture()
    with ZipFile(BytesIO(package)) as source:
        relationships = source.read("word/_rels/document.xml.rels")
    relationships = relationships.replace(
        b"</Relationships>",
        b'<Relationship Id="rExternalLabel" '
        b'Type="http://schemas.example.test/relationships/vbaProject" '
        b'Target="https://127.0.0.1:1/private" TargetMode="External"/>'
        b"</Relationships>",
    )

    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(
            BytesIO(_replace_member(package, "word/_rels/document.xml.rels", relationships))
        )
    assert caught.value.code == "no_extractable_text"


def test_docx_parser_rejects_unicode_normalized_duplicate_part_names() -> None:
    package = _append_member(empty_docx_fixture(), "word/caf\u00e9.xml", b"<x/>")
    package = _append_member(package, "word/cafe\u0301.xml", b"<x/>")
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(package))
    assert caught.value.code == "parser_failed"


def test_docx_parser_rejects_encoded_relationship_target_escape() -> None:
    package = empty_docx_fixture()
    with ZipFile(BytesIO(package)) as source:
        relationships = source.read("word/_rels/document.xml.rels")
    relationships = relationships.replace(
        b"</Relationships>",
        b'<Relationship Id="unsafe" Type="private" Target="../../%2e%2e/private.xml"/>'
        b"</Relationships>",
    )
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(
            BytesIO(
                _replace_member(package, "word/_rels/document.xml.rels", relationships)
            )
        )
    assert caught.value.code == "parser_failed"


@pytest.mark.parametrize(
    "name",
    ["WORD/DOCUMENT.XML", "word/%64ocument.xml", "word/%2e%2e/private.xml"],
)
def test_docx_parser_rejects_canonical_duplicate_or_encoded_hazard_names(name: str) -> None:
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_append_member(empty_docx_fixture(), name, b"<x/>")))
    assert caught.value.code == "parser_failed"


def test_docx_parser_rejects_xml_work_before_document_construction(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from cairn_worker.parsers import docx as docx_parser
    from cairn_worker.parsers import office_safety

    constructed = [False]
    monkeypatch.setattr(office_safety, "OPC_MAX_XML_ELEMENTS", 2)

    def unexpected_document(*args: object, **kwargs: object) -> object:
        del args, kwargs
        constructed[0] = True
        raise AssertionError

    monkeypatch.setattr(docx_parser, "Document", unexpected_document)
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_structured_docx()))
    assert caught.value.code == "parser_failed"
    assert constructed[0] is False


@pytest.mark.parametrize(("kind", "limit"), [("body", 2), ("cells", 1)])
def test_docx_parser_runtime_work_limits_have_exact_boundaries(
    monkeypatch: pytest.MonkeyPatch,
    kind: str,
    limit: int,
) -> None:
    from cairn_worker.parsers import docx as docx_parser

    document = Document()
    if kind == "body":
        document.add_paragraph("one")
        document.add_paragraph("two")
        monkeypatch.setattr(docx_parser, "DOCX_MAX_BODY_ELEMENTS", limit)
        DocxParser().parse(BytesIO(_save_docx(document)))
        document.add_paragraph("three")
    else:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "one"
        monkeypatch.setattr(docx_parser, "DOCX_MAX_TABLE_CELLS", limit)
        DocxParser().parse(BytesIO(_save_docx(document)))
        table.add_row().cells[0].text = "two"
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_save_docx(document)))
    assert caught.value.code == "parser_failed"


def test_docx_style_resolution_rejects_cycles_and_excess_depth_incrementally(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from cairn_worker.parsers import docx as docx_parser

    class Style:
        def __init__(self) -> None:
            self.font = type("Font", (), {"hidden": None})()
            self.base_style: Style | None = None

    first = Style()
    second = Style()
    first.base_style = second
    second.base_style = first
    monkeypatch.setattr(docx_parser, "DOCX_MAX_STYLE_DEPTH", 2, raising=False)
    with pytest.raises(ValueError):
        docx_parser._style_hidden(first)  # pyright: ignore[reportPrivateUsage]

    chain = [Style() for _ in range(4)]
    for parent, child in pairwise(chain):
        parent.base_style = child
    with pytest.raises(ValueError):
        docx_parser._style_hidden(chain[0])  # pyright: ignore[reportPrivateUsage]


@pytest.mark.parametrize("style_kind", ["character", "paragraph"])
def test_docx_parser_rejects_effective_hidden_style_cycles(
    monkeypatch: pytest.MonkeyPatch,
    style_kind: str,
) -> None:
    from cairn_worker.parsers import docx as docx_parser

    document = Document()
    style_type = (
        WD_STYLE_TYPE.CHARACTER
        if style_kind == "character"
        else WD_STYLE_TYPE.PARAGRAPH
    )
    styles = cast(Any, document.styles)
    first = styles.add_style("CycleOne", style_type)
    second = styles.add_style("CycleTwo", style_type)
    first.base_style = second
    second.base_style = first
    if style_kind == "character":
        run = document.add_paragraph().add_run("visible")
        run.style = first
    else:
        document.add_paragraph("visible", style=first)
    monkeypatch.setattr(docx_parser, "DOCX_MAX_STYLE_DEPTH", 4)
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_save_docx(document)))
    assert caught.value.code == "parser_failed"


def test_docx_recursive_revision_walker_owns_visible_wrapped_text_once() -> None:
    document = Document()
    document.add_paragraph("anchor")
    package = _save_docx(document)
    with ZipFile(BytesIO(package)) as source:
        xml = source.read("word/document.xml")
    paragraph = (
        b"<w:p>"
        b"<w:sdt><w:sdtContent><w:r><w:t>sdt</w:t></w:r></w:sdtContent></w:sdt>"
        b"<w:fldSimple w:instr=\"DATE\"><w:r><w:t> field</w:t></w:r></w:fldSimple>"
        b"<w:moveTo><w:r><w:t> moved</w:t></w:r></w:moveTo>"
        b"<w:moveFrom><w:r><w:t> secret</w:t></w:r></w:moveFrom>"
        b"</w:p>"
    )
    xml = xml.replace(b"<w:p><w:r><w:t>anchor</w:t></w:r></w:p>", paragraph)
    blocks = DocxParser().parse(BytesIO(_replace_member(package, "word/document.xml", xml)))
    assert [block.text for block in blocks] == ["sdt field moved"]


def test_docx_recursive_block_walker_preserves_wrapped_table_preorder() -> None:
    package = _nested_merged_hidden_docx()
    with ZipFile(BytesIO(package)) as source:
        xml = source.read("word/document.xml")
    first_table = xml.index(b"<w:tbl>")
    cursor = first_table
    depth = 0
    while True:
        opening = xml.find(b"<w:tbl>", cursor)
        closing = xml.find(b"</w:tbl>", cursor)
        if opening >= 0 and opening < closing:
            depth += 1
            cursor = opening + len(b"<w:tbl>")
        else:
            depth -= 1
            cursor = closing + len(b"</w:tbl>")
            if depth == 0:
                table_end = cursor
                break
    table = xml[first_table:table_end]
    xml = xml[:first_table] + b"<w:sdt><w:sdtContent>" + table + b"</w:sdtContent></w:sdt>" + xml[table_end:]
    blocks = DocxParser().parse(BytesIO(_replace_member(package, "word/document.xml", xml)))
    assert [block.text for block in blocks[:3]] == [
        "Heading",
        "outer\nnested\tright\nhorizontal",
        "vertical",
    ]
    assert [block.locator.model_dump()["table"] for block in blocks[1:3]] == [1, 2]


def test_docx_nested_table_text_is_owned_once_at_exact_cell_source_position() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "before"
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "nested"
    cell.add_paragraph("after")

    blocks = DocxParser().parse(BytesIO(_save_docx(document)))

    assert [block.text for block in blocks] == ["before\nnested\nafter"]
    assert [block.locator.model_dump()["table"] for block in blocks] == [1]
    assert blocks[0].text.count("nested") == 1


def test_docx_table_cell_limit_stops_before_excess_cell_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from cairn_worker.parsers import docx as docx_parser

    document = Document()
    table = document.add_table(rows=1, cols=3)
    for index, cell in enumerate(table.rows[0].cells):
        cell.text = str(index)
    calls = [0]
    original = docx_parser._cell_text  # pyright: ignore[reportPrivateUsage]

    def tracking(cell: object, work: list[int]) -> str:
        calls[0] += 1
        return original(cell, work)  # pyright: ignore[reportArgumentType]

    monkeypatch.setattr(docx_parser, "DOCX_MAX_TABLE_CELLS", 1)
    monkeypatch.setattr(docx_parser, "_cell_text", tracking)
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_save_docx(document)))
    assert caught.value.code == "parser_failed"
    assert calls[0] == 1


@pytest.mark.parametrize(
    "name",
    [
        "word//alias.xml",
        "word/./alias.xml",
        "word/%2Falias.xml",
        "word/%5Calias.xml",
        "word/%2E/alias.xml",
    ],
)
def test_docx_parser_rejects_separator_and_dot_alias_members(name: str) -> None:
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(_append_member(_structured_docx(), name, b"<x/>")))
    assert caught.value.code == "parser_failed"


@pytest.mark.parametrize("target", ["media//x.png", "./media/x.png", "%2Fprivate.xml", "%5Cprivate.xml"])
def test_docx_parser_rejects_ambiguous_relationship_target_separators(target: str) -> None:
    package = _structured_docx()
    with ZipFile(BytesIO(package)) as source:
        relationships = source.read("word/_rels/document.xml.rels")
    relationships = relationships.replace(
        b"</Relationships>",
        f'<Relationship Id="alias" Type="private" Target="{target}"/></Relationships>'.encode(),
    )
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(
            BytesIO(_replace_member(package, "word/_rels/document.xml.rels", relationships))
        )
    assert caught.value.code == "parser_failed"


@pytest.mark.parametrize("damage", ["crc", "overlap"])
def test_docx_parser_validates_all_members_before_document_construction(
    monkeypatch: pytest.MonkeyPatch,
    damage: str,
) -> None:
    from cairn_worker.parsers import docx as docx_parser

    package = _append_member(_structured_docx(), "word/media/blob.bin", b"binary payload")
    package = (
        _corrupt_member_payload(package, "word/media/blob.bin")
        if damage == "crc"
        else _overlap_member_payload(package, "word/media/blob.bin")
    )
    constructed = [False]

    def unexpected_document(*args: object, **kwargs: object) -> object:
        del args, kwargs
        constructed[0] = True
        raise AssertionError

    monkeypatch.setattr(docx_parser, "Document", unexpected_document)
    with pytest.raises(WorkerFailure) as caught:
        DocxParser().parse(BytesIO(package))
    assert caught.value.code == "parser_failed"
    assert constructed[0] is False
